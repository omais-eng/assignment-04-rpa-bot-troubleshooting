"""
Take section screenshots of the RPA website and embed them into Word docs.
Uses Chrome in headless mode via Selenium.
"""
import os, time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By

# ── Paths ────────────────────────────────────────────────────────────────────
BASE   = "/Users/omaissaeedsiddiqui/Desktop/Assignment_04_RPA"
HTML   = f"file://{BASE}/code/index.html"
SS_DIR = f"{BASE}/screenshots"
os.makedirs(SS_DIR, exist_ok=True)

# ── Chrome headless ───────────────────────────────────────────────────────────
opts = Options()
opts.add_argument("--headless=new")
opts.add_argument("--no-sandbox")
opts.add_argument("--disable-dev-shm-usage")
opts.add_argument("--window-size=1400,900")
opts.add_argument("--force-device-scale-factor=1.5")   # crisp screenshots
opts.binary_location = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"

driver = webdriver.Chrome(options=opts)
driver.set_window_size(1400, 900)

try:
    driver.get(HTML)
    time.sleep(2)   # let fonts/CSS settle

    # ── Screenshot definitions ────────────────────────────────────────────────
    sections = [
        ("hero",        None,           "01_Hero_Banner.png"),
        ("errors",      "#errors",      "02_Section1_Errors.png"),
        ("debugging",   "#debugging",   "03_Section2_Debugging.png"),
        ("optimization","#optimization","04_Section3_Optimization.png"),
        ("dashboard",   "#dashboard",   "05_Section4_Dashboard.png"),
        ("ai",          "#ai",          "06_Section5_AI_Usage.png"),
    ]

    for name, anchor, filename in sections:
        if anchor:
            el = driver.find_element(By.CSS_SELECTOR, anchor)
            driver.execute_script("arguments[0].scrollIntoView({block:'start'});", el)
            # extra scroll up so the sticky nav doesn't cut content
            driver.execute_script("window.scrollBy(0, -70);")
        else:
            driver.execute_script("window.scrollTo(0,0);")
        time.sleep(0.8)
        # Expand window height to capture full section
        driver.set_window_size(1400, 900)
        path = f"{SS_DIR}/{filename}"
        driver.save_screenshot(path)
        print(f"  Saved: {filename}")

    # Full-page scroll shots for overview
    driver.execute_script("window.scrollTo(0,0);")
    time.sleep(0.5)
    driver.set_window_size(1400, 900)
    driver.save_screenshot(f"{SS_DIR}/00_Full_Page_Top.png")
    print("  Saved: 00_Full_Page_Top.png")

finally:
    driver.quit()

print("All screenshots saved to:", SS_DIR)

# ── Now embed screenshots into Word docs ──────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY  = RGBColor(0x1a, 0x3a, 0x5c)
GOLD  = RGBColor(0xf5, 0x9e, 0x0b)
BLUE  = RGBColor(0x25, 0x63, 0xeb)
GREEN = RGBColor(0x10, 0xb9, 0x81)
RED   = RGBColor(0xef, 0x44, 0x44)
WHITE = RGBColor(0xff, 0xff, 0xff)
MGRAY = RGBColor(0x64, 0x74, 0x8b)

WORD_DIR = f"{BASE}/word-docs"

def new_doc(title):
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = "Omais Siddiqui"
    sec = doc.sections[0]
    sec.page_width   = Inches(8.5)
    sec.page_height  = Inches(11)
    sec.left_margin  = sec.right_margin  = Inches(0.85)
    sec.top_margin   = sec.bottom_margin = Inches(0.85)
    return doc

def add_title(doc, text, color=NAVY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10); run.font.color.rgb = MGRAY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)

def add_h2(doc, text, color=NAVY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13); run.font.bold = True; run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)

def add_para(doc, text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(4)

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)

def insert_image(doc, img_path, caption="", width=Inches(6.5)):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_picture(img_path, width=width)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph()
            cr = cp.add_run(caption)
            cr.font.size = Pt(9); cr.font.italic = True; cr.font.color.rgb = MGRAY
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
    else:
        add_para(doc, f"[Screenshot not found: {img_path}]")

def page_break(doc):
    doc.add_page_break()

# ── SINGLE COMBINED WORD DOCUMENT ────────────────────────────────────────────
doc1 = new_doc("RPA Assignment 04 — E-Commerce Order Processing RPA Bot")
add_title(doc1, "E-Commerce Order Processing RPA Bot")
add_title(doc1, "Troubleshooting & Debugging Case Study")
add_subtitle(doc1, "RPA Course  |  Assignment 04  |  Omais Siddiqui  |  April 2026")

# ── TABLE OF CONTENTS ─────────────────────────────────────────────────────────
add_h2(doc1, "Contents")
for item in [
    "Section 1 — RPA Error Log: 5 Realistic Exceptions",
    "Section 2 — Debugging Strategy & Framework",
    "Section 3 — Bot Optimization Strategies",
    "Section 4 — Monitoring Dashboard",
    "Section 5 — AI-Assisted Development (Vibe Coding)",
]:
    add_bullet(doc1, item)

page_break(doc1)

# ── SECTION 1 ─────────────────────────────────────────────────────────────────
add_title(doc1, "Section 1 — RPA Error Log: 5 Realistic Exceptions", color=NAVY)
add_subtitle(doc1, "Documented errors across workflow stages categorized by type, impact, and root cause")

# ── SECTION 1 CONTENT ────────────────────────────────────────────────────────
add_h2(doc1, "Website Screenshot — Hero & Error Log")
insert_image(doc1, f"{SS_DIR}/01_Hero_Banner.png",       "Figure 1: Website Hero Banner")
insert_image(doc1, f"{SS_DIR}/02_Section1_Errors.png",   "Figure 2: Section 1 — RPA Error Log Table")

page_break(doc1)
add_h2(doc1, "5 Realistic RPA Exceptions")
errors = [
    ("E-01", "SelectorNotFoundException", "System",   "Order Placement",    "HIGH",
     "Checkout portal frontend renamed CSS selector from #btn-checkout to .submit-order-cta during a release. "
     "The bot's selector library was not updated. This blocked approximately 120 orders per hour."),
    ("E-02", "InvalidOrderDataException", "Business", "Data Extraction",    "HIGH",
     "ERP system began exporting international orders without the postal code field. "
     "Input schema validation layer never updated, causing silent shipment failures downstream."),
    ("E-03", "OracleDBTimeoutException",  "System",   "Inventory Check",    "HIGH",
     "Inventory table lacks composite index on (product_sku, warehouse_id). "
     "Full table scans under 50+ concurrent bots caused 30s timeouts — 400+ units oversold per week."),
    ("E-04", "DuplicateOrderException",   "Business", "Order Submission",   "MEDIUM",
     "Retry logic lacked idempotency key check. On network timeout, bot retried without verifying "
     "if original POST succeeded — creating duplicate records and double-charging customers."),
    ("E-05", "PDFParsingFailureException","App",      "Invoice Generation", "LOW",
     "Vendor switched invoice template from text-layer PDF to scanned image PDF. "
     "OCR pre-processing not enabled in PDFActivityPack — parser returned null fields."),
]
for eid, name, etype, stage, sev, desc in errors:
    add_h2(doc1, f"{eid} — {name}", color=BLUE)
    for label, val in [("Type", etype), ("Stage", stage), ("Severity", sev)]:
        p = doc1.add_paragraph()
        r1 = p.add_run(f"{label}: "); r1.font.bold = True; r1.font.size = Pt(10.5)
        r2 = p.add_run(val); r2.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)
    add_para(doc1, f"Root Cause: {desc}")

# ── SECTION 2 ─────────────────────────────────────────────────────────────────
page_break(doc1)
add_title(doc1, "Section 2 — Debugging Strategy & Framework", color=NAVY)
add_subtitle(doc1, "Logging, monitoring, reproduction steps, retry logic, and decision flow")

add_h2(doc1, "Website Screenshot — Debugging Strategy")
insert_image(doc1, f"{SS_DIR}/03_Section2_Debugging.png", "Figure 3: Section 2 — Debugging Strategy & Decision Flow")

page_break(doc1)
add_h2(doc1, "1. Logging Strategy")
for bullet in [
    "Structured levels: INFO (normal flow), WARN (recoverable anomalies), ERROR (exceptions + stack traces), FATAL (process-terminating).",
    "Every log entry includes: OrderID, BotInstanceID, UTC Timestamp, WorkflowStage, CorrelationID.",
    "Log shipping pipeline: UiPath Orchestrator → Elasticsearch → Kibana. Retention: 90 days.",
    "Screenshot on every exception — saved with OrderID-Timestamp naming convention.",
]:
    add_bullet(doc1, bullet)

add_h2(doc1, "2. Monitoring & Alerting")
for bullet in [
    "Real-time KPI tracking: success rate, exception rate, avg processing time, queue depth — 60s refresh.",
    "PagerDuty alerts: success rate <90%, error rate >5%, queue depth >500 items, bot unresponsive >5 min.",
    "SLA violation detection: orders not processed within 4 hours trigger email to Operations Manager.",
    "Anomaly detection: unusual error spikes auto-create Jira tickets with log snippets.",
]:
    add_bullet(doc1, bullet)

add_h2(doc1, "3. Error Reproduction Steps")
for i, step in enumerate([
    "Isolate: Extract the failing transaction from Orchestrator queue with exact input payload.",
    "Snapshot: Clone production environment state (app version, DB snapshot, selector library).",
    "Replay: Run bot in UiPath Debug mode against the isolated transaction.",
    "Observe: Use Step Into / Step Over; inspect variables in Watch panel.",
    "Confirm: Reproduce consistently at least 3× before applying any fix.",
    "Validate: Run 50 representative test transactions — zero failures required for production promotion.",
], 1):
    p = doc1.add_paragraph()
    r1 = p.add_run(f"Step {i}: "); r1.font.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = BLUE
    r2 = p.add_run(step); r2.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.2)

add_h2(doc1, "4. Retry Logic & Exception Handling")
for exc, action in [
    ("SelectorNotFoundException", "Refresh selector library from Orchestrator asset → Send alert → Rethrow."),
    ("InvalidOrderDataException", "Move to Exception Queue → Log warning → Continue next order."),
    ("TimeoutException",          "Wait 10s → Retry. If retries exhausted → Escalate to human queue."),
    ("Finally block",             "Always: take screenshot, close applications, release DB connections."),
]:
    p = doc1.add_paragraph()
    r1 = p.add_run(f"• {exc}: "); r1.font.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = RED
    r2 = p.add_run(action); r2.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.2)

add_h2(doc1, "5. Debugging Decision Flow (10 Steps)")
for i, (step, desc) in enumerate([
    ("Error Detected",   "Bot throws exception or enters Catch block."),
    ("Log Captured",     "Structured log entry written with full context data."),
    ("Alert Triggered",  "PagerDuty/email notification sent based on severity."),
    ("Classify Type",    "Determine: System, Business, or Application exception."),
    ("Reproduce in DEV", "Isolate transaction, snapshot environment, replay in debug."),
    ("Apply Fix",        "Code fix, config update, or selector library update in Dev."),
    ("Test 50 Orders",   "Run 50 representative transactions — zero failures required."),
    ("Deploy to PROD",   "Merge to main, deploy via Orchestrator package manager."),
    ("Monitor 24hr",     "Watch KPI dashboard actively for 24 hours post-deployment."),
    ("Close Incident",   "Log root cause analysis, update runbook, close Jira ticket."),
], 1):
    p = doc1.add_paragraph()
    r1 = p.add_run(f"{i:02d}. {step}: "); r1.font.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = NAVY
    r2 = p.add_run(desc); r2.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.2)

# ── SECTION 3 ─────────────────────────────────────────────────────────────────
page_break(doc1)
add_title(doc1, "Section 3 — Bot Optimization Strategies", color=NAVY)
add_subtitle(doc1, "Performance, reliability, maintainability, scheduling & before/after comparison")

add_h2(doc1, "Website Screenshot — Optimization Section")
insert_image(doc1, f"{SS_DIR}/04_Section3_Optimization.png", "Figure 4: Section 3 — Optimization Strategies")

page_break(doc1)
for section_title, color, items in [
    ("Performance Improvements", BLUE, [
        "Parallel processing: 5 concurrent bot lanes via Work Queues — 5× throughput.",
        "DB composite index on (product_sku, warehouse_id): 30s timeout → <200ms.",
        "Lazy UI loading (WaitForReady=Complete only on critical elements): 40% wait reduction.",
        "Batch inventory API calls in groups of 50 instead of individual requests.",
        "Connection pooling: reuse DB connections across transactions.",
    ]),
    ("Reliability Improvements", GREEN, [
        "Idempotency keys (UUID per order) — eliminates all duplicate orders (fixes E-04).",
        "Dynamic anchor-based selectors with fuzzy matching — survives UI updates (fixes E-01).",
        "JSON schema validation on ERP feed before processing begins (fixes E-02).",
        "Circuit breaker: if DB error rate >10% in 60s, pause bot and alert.",
        "30-second heartbeat pings to Orchestrator for rapid crash detection.",
    ]),
    ("Maintainability Improvements", RGBColor(0xf5,0x9e,0x0b), [
        "Centralized selector library in Orchestrator assets — single update point.",
        "Modular .xaml workflows per stage with defined inputs/outputs.",
        "All parameters (batch sizes, timeouts, retries) in Orchestrator assets — zero code deploys.",
        "OCR pre-processing layer auto-detects scanned PDFs → Tesseract OCR (fixes E-05).",
        "All changes require peer review + automated regression test suite.",
    ]),
    ("Scheduling Strategy", RGBColor(0x93,0x33,0xea), [
        "Peak hours 8AM–6PM: 5 bot instances — maximum throughput.",
        "Off-peak 6PM–8AM: 2 instances — process backlog at lower cost.",
        "Maintenance window 2AM–4AM daily: deployments, index rebuilds, selector updates.",
        "VIP/same-day orders jump queue: 15-minute SLA guarantee.",
        "Weekend: 3 instances with auto scale-up if queue exceeds 200 items.",
    ]),
]:
    add_h2(doc1, section_title, color=color)
    for item in items:
        add_bullet(doc1, item)

add_h2(doc1, "Before vs. After — Key Metrics")
for metric, before, after in [
    ("Success Rate",      "71%",                 "97.4% (+26%)"),
    ("Processing Time",   "4.2 min/order",       "1.8 min/order (57% faster)"),
    ("Throughput",        "120 orders/hour",      "600 orders/hour (5× increase)"),
    ("Duplicate Orders",  "~12 incidents/week",  "0 incidents/month"),
    ("DB Query Time",     "30s+ (timeout)",       "<200ms (indexed)"),
    ("Manual Overhead",   "3.5 hours/day",        "12 minutes/day (96% reduction)"),
    ("Selector Failures", "After every deploy",   "Zero — dynamic selectors"),
    ("Config Changes",    "Full code deployment", "Zero-downtime via Orchestrator"),
]:
    p = doc1.add_paragraph()
    r1 = p.add_run(f"• {metric}: "); r1.font.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(f"{before}  →  "); r2.font.size = Pt(10.5); r2.font.color.rgb = RED
    r3 = p.add_run(after); r3.font.size = Pt(10.5); r3.font.color.rgb = GREEN
    p.paragraph_format.left_indent = Inches(0.2)

# ── SECTION 4 ─────────────────────────────────────────────────────────────────
page_break(doc1)
add_title(doc1, "Section 4 — Monitoring Dashboard", color=NAVY)
add_subtitle(doc1, "KPIs, queue health, alerts, and bot instance status")

add_h2(doc1, "Website Screenshot — Monitoring Dashboard")
insert_image(doc1, f"{SS_DIR}/05_Section4_Dashboard.png", "Figure 5: Section 4 — Monitoring Dashboard Mock UI with KPIs")

page_break(doc1)
add_h2(doc1, "Dashboard KPIs Reference")
for kpi, value, target, status in [
    ("Total Orders Today",   "4,827",     "N/A — volume metric", "Normal"),
    ("Success Rate",         "97.4%",     "≥ 95%",               "Exceeds SLA"),
    ("Failed Orders",        "125 (2.6%)","≤ 5%",                "Within Range"),
    ("Retries Executed",     "218 (4.5%)","≤ 10%",               "Normal"),
    ("Avg Processing Time",  "1.8 min",   "≤ 4 min",             "Exceeds SLA"),
    ("Bot Uptime (30 days)", "99.2%",     "≥ 99%",               "Exceeds SLA"),
    ("Queue Depth (New)",    "47 items",  "< 200 items",         "Normal"),
    ("Exception Queue",      "23 items",  "< 50 items",          "Monitor"),
    ("Dead Letter Queue",    "5 items",   "0 items",             "Action Required"),
]:
    p = doc1.add_paragraph()
    r1 = p.add_run(f"• {kpi}: "); r1.font.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(f"{value} | Target: {target} | Status: {status}"); r2.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.2)

# ── SECTION 5 ─────────────────────────────────────────────────────────────────
page_break(doc1)
add_title(doc1, "Section 5 — AI-Assisted Development (Vibe Coding)", color=NAVY)
add_subtitle(doc1, "How Claude Code was used to build this project")

add_h2(doc1, "Website Screenshot — AI Usage Section")
insert_image(doc1, f"{SS_DIR}/06_Section5_AI_Usage.png", "Figure 6: Section 5 — AI-Assisted Development with Claude Code")

page_break(doc1)
add_h2(doc1, "AI-Assisted Development: Vibe Coding with Claude Code")
add_para(doc1,
    "Vibe coding is a modern AI-assisted development approach where the developer describes what they want "
    "in natural language, and Claude Code generates, refines, and iterates on the actual code. "
    "The developer acts as architect and reviewer — the AI handles implementation details.", 10.5)
for title, detail in [
    ("Prompt Engineering",  "Full project requirements described in natural language → Claude structured architecture, layout, and data models."),
    ("UI/UX Generation",    "Design specs like 'enterprise-style, navy/gold color scheme' became actual CSS gradients and grid systems."),
    ("Domain Synthesis",    "RPA error scenarios with realistic names and root causes synthesized from domain context."),
    ("Iterative Refinement","Prompt → generate → review → refine loop. The decision flow evolved from a list to a visual step flow."),
    ("Full Stack Generation","Generated complete HTML/CSS/JS website, GitHub repository setup, and video script — all through an AI-assisted development workflow."),
    ("Consistent Mock Data","Dashboard KPIs internally consistent; before/after metrics logically aligned with optimizations."),
]:
    p = doc1.add_paragraph()
    r1 = p.add_run(f"• {title}: "); r1.font.bold = True; r1.font.size = Pt(10.5)
    r2 = p.add_run(detail); r2.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)

# ── SAVE SINGLE DOCUMENT ──────────────────────────────────────────────────────
out_path = f"{WORD_DIR}/Assignment_04_RPA_Complete.docx"
doc1.save(out_path)
print(f"Saved: {out_path}")
print("\n✅ Single Word document generated with all sections and screenshots.")
