from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

NAVY  = RGBColor(0x1a, 0x3a, 0x5c)
GOLD  = RGBColor(0xf5, 0x9e, 0x0b)
BLUE  = RGBColor(0x25, 0x63, 0xeb)
GREEN = RGBColor(0x10, 0xb9, 0x81)
RED   = RGBColor(0xef, 0x44, 0x44)
WHITE = RGBColor(0xff, 0xff, 0xff)
MGRAY = RGBColor(0x64, 0x74, 0x8b)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=10, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1, color=NAVY):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(18 if level==1 else 14 if level==2 else 12)
    run.font.bold = True
    run.font.color.rgb = color
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(6)
    return para

def add_para(doc, text, size=10.5, color=None, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = color
    para.paragraph_format.space_after = Pt(4)
    return para

def add_bullet(doc, text, size=10):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(2)

# ──────────────────────────────────────────────
# DOCUMENT 1 — ERRORS
# ──────────────────────────────────────────────
doc1 = Document()
doc1.core_properties.title = "RPA Assignment 04 - Section 1: Error Log"
doc1.core_properties.author = "Omais Siddiqui"
sec = doc1.sections[0]
sec.page_width  = Inches(11)
sec.page_height = Inches(8.5)
sec.left_margin = sec.right_margin = Inches(0.8)
sec.top_margin  = sec.bottom_margin = Inches(0.7)

p = doc1.add_paragraph()
run = p.add_run("RPA Bot Troubleshooting — Section 1: Error Log")
run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)

p2 = doc1.add_paragraph()
run2 = p2.add_run("Assignment 04  |  E-Commerce Order Processing RPA Bot  |  Omais Siddiqui  |  April 2026")
run2.font.size = Pt(10); run2.font.color.rgb = MGRAY
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(12)

add_heading(doc1, "5 Realistic RPA Exceptions — System, Business & Application", 2)
add_para(doc1, "The following table documents five realistic exceptions encountered by the E-Commerce Order Processing RPA Bot. Each error is categorized by type, workflow stage, business impact, and root cause.", size=10)

headers = ["#", "Error Name", "Type", "Stage", "Impact", "Root Cause", "Severity"]
errors = [
    ("E-01", "SelectorNotFoundException",    "System",   "Order Placement",    "~120 orders/hr blocked; orders never submitted",                                              "Checkout portal frontend renamed CSS selector from #btn-checkout to .submit-order-cta during a release. Selector library not updated.",            "HIGH"),
    ("E-02", "InvalidOrderDataException",    "Business", "Data Extraction",    "Orders with missing ZIP codes cause failed shipments and customer complaints",               "ERP system began exporting international orders without postal code field. Input schema validation layer never updated.",                            "HIGH"),
    ("E-03", "OracleDBTimeoutException",     "System",   "Inventory Check",    "400+ units oversold per week; orders proceed without stock confirmation",                    "Inventory table lacks composite index on product_sku + warehouse_id. Full table scans under 50+ concurrent bots cause 30s timeouts.",               "HIGH"),
    ("E-04", "DuplicateOrderException",      "Business", "Order Submission",   "Customers charged twice; 2-3 hrs/incident manual correction in finance",                     "Retry logic lacked idempotency key check. On network timeout, bot retried without checking if original POST succeeded — creating duplicates.",       "MEDIUM"),
    ("E-05", "PDFParsingFailureException",   "App",      "Invoice Generation", "15% of daily invoices require manual re-generation; 1.5 hr/day overhead",                   "Vendor switched invoice template from text-layer PDF to scanned image PDF. OCR pre-processing not enabled in PDFActivityPack.",                        "LOW"),
]
col_widths = [Inches(0.45), Inches(1.9), Inches(0.75), Inches(1.1), Inches(1.8), Inches(2.8), Inches(0.75)]
table = doc1.add_table(rows=1, cols=7)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for j, (h, cell) in enumerate(zip(headers, hdr_cells)):
    set_cell_bg(cell, "1a3a5c")
    cell_text(cell, h, bold=True, size=9.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell.width = col_widths[j]

type_hex  = {"System": "dbeafe", "Business": "fef3c7", "App": "ede9fe"}
type_txt  = {"System": BLUE,     "Business": RGBColor(0x92,0x40,0x0e), "App": RGBColor(0x5b,0x21,0xb6)}
sev_hex   = {"HIGH": "fee2e2", "MEDIUM": "fef3c7", "LOW": "d1fae5"}
sev_txt   = {"HIGH": RED,      "MEDIUM": RGBColor(0x92,0x40,0x0e), "LOW": GREEN}
row_hex   = ["ffffff","f8fafc","ffffff","f8fafc","ffffff"]

for i, (eid, name, etype, stage, impact, cause, sev) in enumerate(errors):
    row = table.add_row().cells
    set_cell_bg(row[0], row_hex[i]); cell_text(row[0], eid,    bold=True,  size=9, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row[1], row_hex[i]); cell_text(row[1], name,   bold=True,  size=9, color=NAVY)
    set_cell_bg(row[2], type_hex[etype]); cell_text(row[2], etype, bold=True, size=9, color=type_txt[etype], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row[3], row_hex[i]); cell_text(row[3], stage,  size=9)
    set_cell_bg(row[4], row_hex[i]); cell_text(row[4], impact, size=9)
    set_cell_bg(row[5], row_hex[i]); cell_text(row[5], cause,  size=9)
    set_cell_bg(row[6], sev_hex[sev]); cell_text(row[6], sev, bold=True, size=9, color=sev_txt[sev], align=WD_ALIGN_PARAGRAPH.CENTER)
    for j in range(7): row[j].width = col_widths[j]

doc1.add_paragraph()
add_heading(doc1, "Error Type Distribution", 2)
for label, ids, detail in [
    ("System Exceptions (2)", "E-01, E-03", "Infrastructure failures: stale UI selector and database query timeout under concurrency."),
    ("Business Exceptions (2)", "E-02, E-04", "Data and logic failures: invalid input data and missing idempotency in retry logic."),
    ("Application Exceptions (1)", "E-05", "Parser failure: vendor PDF format change from text-layer to rasterized scanned image."),
]:
    add_bullet(doc1, f"{label} — {ids}: {detail}")

out1 = "/Users/omaissaeedsiddiqui/Desktop/Assignment_04_RPA/word-docs/Section_1_Error_Log.docx"
doc1.save(out1)
print(f"Saved: {out1}")

# ──────────────────────────────────────────────
# DOCUMENT 2 — DEBUGGING STRATEGY
# ──────────────────────────────────────────────
doc2 = Document()
doc2.core_properties.title = "RPA Assignment 04 - Section 2: Debugging Strategy"
doc2.core_properties.author = "Omais Siddiqui"
sec = doc2.sections[0]
sec.left_margin = sec.right_margin = Inches(1)
sec.top_margin  = sec.bottom_margin = Inches(0.9)

p = doc2.add_paragraph()
run = p.add_run("Section 2 — Debugging Strategy & Framework")
run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc2, "1. Logging Strategy", 2)
items = [
    ("Structured Log Levels", "Use INFO for normal flow milestones, WARN for recoverable anomalies, ERROR for exceptions with full stack traces, FATAL for process-terminating failures."),
    ("Contextual Data", "Every log entry must include: OrderID, BotInstanceID, Timestamp (UTC), WorkflowStage, CorrelationID for cross-service tracing."),
    ("Log Shipping Pipeline", "UiPath Orchestrator → Elasticsearch → Kibana. Log retention: 90 days for audit compliance."),
    ("Screenshot on Exception", "Take Screenshot activity triggered in every Catch block, saved to shared drive with OrderID-Timestamp naming."),
]
for title, detail in items:
    para = doc2.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(f"• {title}: "); run.font.bold = True; run.font.size = Pt(10.5)
    run2 = para.add_run(detail); run2.font.size = Pt(10.5)

add_heading(doc2, "2. Monitoring & Alerting", 2)
items2 = [
    ("Real-Time KPIs", "Success rate, exception rate, avg processing time, queue depth — refreshed every 60 seconds in Orchestrator."),
    ("PagerDuty Thresholds", "Alert triggers: success rate <90%, error rate >5%, queue depth >500 items, bot unresponsive >5 min."),
    ("SLA Violation Detection", "Orders not processed within 4 hours flagged automatically with email to Operations Manager."),
    ("Anomaly Detection", "Unusual error type spikes trigger automated Jira tickets with log snippets attached."),
]
for title, detail in items2:
    para = doc2.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(f"• {title}: "); run.font.bold = True; run.font.size = Pt(10.5)
    run2 = para.add_run(detail); run2.font.size = Pt(10.5)

add_heading(doc2, "3. Error Reproduction Steps", 2)
steps = [
    "Isolate: Extract the failing transaction from Orchestrator queue with its exact input payload.",
    "Snapshot: Clone production environment state (app version, DB snapshot, selector library version).",
    "Replay: Run bot in Debug mode against the isolated transaction in the test environment.",
    "Observe: Use UiPath's Step Into / Step Over to walk through each activity; inspect variables in Watch panel.",
    "Confirm: Reproduce consistently at least 3× before applying any fix to avoid false positives.",
    "Validate Fix: Run 50 representative test transactions after fix. Zero new failures = approved for production.",
]
for i, step in enumerate(steps, 1):
    para = doc2.add_paragraph()
    run = para.add_run(f"Step {i}: "); run.font.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = BLUE
    run2 = para.add_run(step); run2.font.size = Pt(10.5)
    para.paragraph_format.left_indent = Inches(0.25)

add_heading(doc2, "4. Retry Logic & Exception Handling", 2)
add_para(doc2, "The RetryScope activity is configured with 3 retries and a 5-second exponential backoff. An idempotency key check occurs before every order submission to prevent duplicate processing. Each catch block handles its specific exception type:", size=10.5)
catch_items = [
    ("SelectorNotFoundException", "Refresh selector library from Orchestrator asset, send alert, rethrow."),
    ("InvalidOrderDataException", "Move transaction to Exception Queue, log warning, continue processing next order."),
    ("TimeoutException", "Wait 10 seconds and retry. If retries exhausted, escalate to human intervention queue."),
    ("Finally block", "Always take screenshot, close all applications, release DB connections."),
]
for exc, action in catch_items:
    para = doc2.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    run = para.add_run(f"• {exc}: "); run.font.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = RED
    run2 = para.add_run(action); run2.font.size = Pt(10.5)

add_heading(doc2, "5. Debugging Decision Flow (10 Steps)", 2)
flow_steps = [
    ("Error Detected",   "Bot throws an unhandled exception or enters a Catch block."),
    ("Log Captured",     "Structured log entry written with full context data."),
    ("Alert Triggered",  "PagerDuty/email notification sent based on severity threshold."),
    ("Classify Type",    "Determine: System, Business, or Application exception."),
    ("Reproduce in DEV", "Isolate transaction, snapshot environment, replay in debug mode."),
    ("Apply Fix",        "Code fix, config update, or selector library update deployed to Dev."),
    ("Test 50 Orders",   "Run 50 representative transactions. Zero failures required."),
    ("Deploy to PROD",   "Merge to main branch, deploy via Orchestrator package manager."),
    ("Monitor 24hr",     "Actively watch KPI dashboard for 24 hours post-deployment."),
    ("Close Incident",   "Log root cause analysis, update runbook, close Jira ticket."),
]
for i, (step, desc) in enumerate(flow_steps, 1):
    para = doc2.add_paragraph()
    run = para.add_run(f"{i:02d}. {step}: "); run.font.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = NAVY
    run2 = para.add_run(desc); run2.font.size = Pt(10.5)
    para.paragraph_format.left_indent = Inches(0.25)

out2 = "/Users/omaissaeedsiddiqui/Desktop/Assignment_04_RPA/word-docs/Section_2_Debugging_Strategy.docx"
doc2.save(out2)
print(f"Saved: {out2}")

# ──────────────────────────────────────────────
# DOCUMENT 3 — OPTIMIZATION
# ──────────────────────────────────────────────
doc3 = Document()
doc3.core_properties.title = "RPA Assignment 04 - Section 3: Optimization"
doc3.core_properties.author = "Omais Siddiqui"
sec = doc3.sections[0]
sec.left_margin = sec.right_margin = Inches(1)
sec.top_margin  = sec.bottom_margin = Inches(0.9)

p = doc3.add_paragraph()
run = p.add_run("Section 3 — Bot Optimization Strategies")
run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

opt_sections = [
    ("Performance Improvements", BLUE, [
        "Parallel processing: split order queue into 5 concurrent bot lanes using Work Queues — 5× throughput increase.",
        "Database query optimization: composite index on (product_sku, warehouse_id) — query time 30s → <200ms.",
        "Lazy UI loading: use WaitForReady=Complete only on critical elements — 40% reduction in average wait time.",
        "Batch API calls: group inventory checks into batches of 50 instead of individual requests.",
        "Connection pooling: reuse database connections across transactions.",
    ]),
    ("Reliability Improvements", GREEN, [
        "Idempotency keys: generate UUID per order before submission — eliminates all duplicate orders (fixes E-04).",
        "Dynamic selectors: anchor-based selectors with fuzzy matching survive UI updates (fixes E-01).",
        "Input schema validation: strict JSON schema validation on ERP feed before processing (fixes E-02).",
        "Circuit breaker: if DB error rate >10% in 60s, pause bot and alert — prevents cascading failures.",
        "Heartbeat monitoring: 30-second pings to Orchestrator for rapid crash detection.",
    ]),
    ("Maintainability Improvements", RGBColor(0xf5,0x9e,0x0b), [
        "Centralized selector library: all selectors stored in Orchestrator config assets — single update point.",
        "Modular workflow design: each stage is an independent .xaml file with defined inputs/outputs.",
        "Configuration-driven parameters: batch sizes, timeouts, retry counts in Orchestrator assets — zero code deployments.",
        "OCR pipeline: pre-processing layer auto-detects scanned PDFs and routes to Tesseract OCR (fixes E-05).",
        "Code review checklist: all changes require peer review + automated regression test suite.",
    ]),
    ("Scheduling Strategy", RGBColor(0x93,0x33,0xea), [
        "Peak hours (8AM–6PM): 5 bot instances — maximum throughput mode.",
        "Off-peak (6PM–8AM): 2 bot instances — process backlog at lower resource cost.",
        "Maintenance window: 2AM–4AM daily — selector updates, DB index rebuilds, version deployments.",
        "Priority queue: VIP/same-day orders jump the main queue, processed within 15 minutes.",
        "Weekend: 3 instances with auto scale-up if queue exceeds 200 items.",
    ]),
]
for title, color, items in opt_sections:
    add_heading(doc3, title, 2, color=color)
    for item in items:
        add_bullet(doc3, item)

add_heading(doc3, "Before vs. After Optimization Comparison", 2)
table = doc3.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdrs = ["Metric", "Before", "After"]
hdr_clrs = ["1a3a5c", "7f1d1d", "065f46"]
for j, (h, hc) in enumerate(zip(hdrs, hdr_clrs)):
    set_cell_bg(table.rows[0].cells[j], hc)
    cell_text(table.rows[0].cells[j], h, bold=True, size=11, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

comparisons = [
    ("Success Rate",          "71%",                  "97.4% (+26%)"),
    ("Processing Time",       "4.2 min/order",        "1.8 min/order (57% faster)"),
    ("Throughput",            "120 orders/hour",      "600 orders/hour (5× increase)"),
    ("Duplicate Orders",      "~12 incidents/week",   "0 incidents/month"),
    ("DB Query Time",         "30+ seconds (timeout)","<200ms (indexed)"),
    ("Manual Overhead",       "3.5 hours/day",        "12 minutes/day (96% reduction)"),
    ("Selector Failures",     "After every deploy",   "Zero — dynamic selectors"),
    ("Config Changes",        "Require code deploy",  "Zero-downtime via Orchestrator"),
]
for metric, before, after in comparisons:
    row = table.add_row().cells
    set_cell_bg(row[0], "f8fafc"); cell_text(row[0], metric, bold=True, size=10)
    set_cell_bg(row[1], "fff5f5"); cell_text(row[1], before,  size=10, color=RED)
    set_cell_bg(row[2], "f0fdf4"); cell_text(row[2], after,   size=10, color=GREEN)

out3 = "/Users/omaissaeedsiddiqui/Desktop/Assignment_04_RPA/word-docs/Section_3_Optimization.docx"
doc3.save(out3)
print(f"Saved: {out3}")

# ──────────────────────────────────────────────
# DOCUMENT 4 — DASHBOARD + AI
# ──────────────────────────────────────────────
doc4 = Document()
doc4.core_properties.title = "RPA Assignment 04 - Sections 4 & 5: Dashboard & AI Usage"
doc4.core_properties.author = "Omais Siddiqui"
sec = doc4.sections[0]
sec.left_margin = sec.right_margin = Inches(1)
sec.top_margin  = sec.bottom_margin = Inches(0.9)

p = doc4.add_paragraph()
run = p.add_run("Section 4 — Monitoring Dashboard & Section 5 — AI Usage")
run.font.size = Pt(20); run.font.bold = True; run.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading(doc4, "Section 4: Monitoring Dashboard KPIs", 2)
add_para(doc4, "The following table shows the key performance indicators tracked on the UiPath Orchestrator monitoring dashboard post-optimization.", 10.5)

table = doc4.add_table(rows=1, cols=4)
table.style = 'Table Grid'
for j, h in enumerate(["KPI", "Value", "Target / SLA", "Status"]):
    set_cell_bg(table.rows[0].cells[j], "1a3a5c")
    cell_text(table.rows[0].cells[j], h, bold=True, size=10, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
kpis = [
    ("Total Orders Today",     "4,827",    "N/A — volume metric",  "✅ Normal",  "d1fae5"),
    ("Success Rate",           "97.4%",    "≥ 95%",                "✅ Exceeds", "d1fae5"),
    ("Failed Orders",          "125 (2.6%)","≤ 5%",                "✅ In Range","d1fae5"),
    ("Retries Executed",       "218 (4.5%)","≤ 10%",               "✅ Normal",  "d1fae5"),
    ("Avg Processing Time",    "1.8 min",  "≤ 4 min",              "✅ Exceeds", "d1fae5"),
    ("Bot Uptime (30 days)",   "99.2%",    "≥ 99%",                "✅ Exceeds", "d1fae5"),
    ("Queue Depth (New)",      "47 items", "< 200 items",          "✅ Normal",  "d1fae5"),
    ("Exception Queue",        "23 items", "< 50 items",           "⚠️ Monitor", "fef3c7"),
    ("Dead Letter Queue",      "5 items",  "0 items",              "🚨 Action",  "fee2e2"),
]
for metric, value, target, status, bg in kpis:
    row = table.add_row().cells
    set_cell_bg(row[0], "f8fafc"); cell_text(row[0], metric, bold=True, size=10)
    set_cell_bg(row[1], "f8fafc"); cell_text(row[1], value, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(row[2], "f8fafc"); cell_text(row[2], target, size=10)
    set_cell_bg(row[3], bg);       cell_text(row[3], status, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc4.add_paragraph()
add_heading(doc4, "Active Alerts Summary", 2)
alerts = [
    ("✅ CLEAR",   "System Health — All 5 bots operational. Success rate 97.4% above SLA."),
    ("⚠️ MEDIUM",  "Exception Queue — 23 items, 7 approaching 2-hour SLA deadline."),
    ("ℹ️ INFO",    "Maintenance Tonight — Selector library v2.4 + DB index rebuild at 2:00 AM."),
    ("🚨 CRITICAL","Dead Letter Queue — 5 orders with undocumented error code 9042 from payment gateway."),
]
for severity, detail in alerts:
    para = doc4.add_paragraph()
    run = para.add_run(f"{severity}: "); run.font.bold = True; run.font.size = Pt(10.5)
    run2 = para.add_run(detail); run2.font.size = Pt(10.5)
    para.paragraph_format.left_indent = Inches(0.25)

doc4.add_paragraph()
add_heading(doc4, "Section 5: AI-Assisted Development — Vibe Coding with Claude Code", 1)
add_para(doc4,
    "Vibe coding is a modern AI-assisted development approach where the developer describes what they want in natural language, "
    "and an AI (Claude Code) generates, refines, and iterates on the actual code. The developer acts as architect and reviewer — "
    "the AI handles implementation details.", 10.5)

ai_aspects = [
    ("Prompt Engineering for Architecture",
     "The full project requirements were described in natural language. Claude Code structured the entire component architecture, section layout, and data models from plain English descriptions."),
    ("UI/UX Generation via Natural Language",
     "Design specifications like 'enterprise-style, navy/gold color scheme, card-based layout' were translated into actual CSS custom properties, gradient definitions, and grid systems."),
    ("Domain Content Synthesis",
     "RPA error scenarios with realistic names, root causes, and business impacts were synthesized by feeding Claude the domain context — combining UiPath platform knowledge with e-commerce workflows."),
    ("Iterative Refinement Loop",
     "The development process was iterative: prompt → generate → review → refine. Each section was improved through natural language feedback rather than manual code editing."),
    ("Full Stack Generation",
     "Claude Code generated: HTML/CSS/JS website, Python PPT/Word scripts, folder structure, GitHub setup commands, and the video script — all from conversational prompts."),
    ("Consistent Mock Data",
     "Dashboard KPIs were generated to be internally consistent and realistic. Before/after metrics logically aligned with the described optimization improvements."),
]
for title, detail in ai_aspects:
    add_heading(doc4, title, 3, color=BLUE)
    add_para(doc4, detail, 10.5)

out4 = "/Users/omaissaeedsiddiqui/Desktop/Assignment_04_RPA/word-docs/Sections_4_5_Dashboard_AI.docx"
doc4.save(out4)
print(f"Saved: {out4}")
print("All Word documents generated successfully.")
